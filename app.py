from flask import Flask, request, jsonify
from docx import Document
import os
from flask_cors import CORS
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

app = Flask(__name__)


CORS(app, supports_credentials=True, origins=["https://e-verification-bkfr.vercel.app", "http://127.0.0.1:5000", "http://127.0.0.1:5173", "https://www.bioentrust.net"])

@app.route('/')
def index():
    return 'Welcome to the e-affidavit server!'


@app.route('/generate_affidavit', methods=['POST'])
def generate_affidavit():
    # Get parameters from the request
    first_name = request.json['first_name']
    middle_name = request.json['middle_name']
    last_name = request.json['last_name']
    reason = request.json['reason']
    address = request.json['address']
    religion = request.json['religion']
    mobile = request.json['mobile']
    email = request.json['email']
    signature = request.json['signature']
    date_of_loss = request.json['date_of_loss']
    network = request.json['network']
    date = request.json['date']
    position = request.json['position']
    gender = request.json['gender']
    company_address = request.json['company_address']
    seal_sign = request.json['seal_sign']


    # Open the DOCX file in the project directory
    docx_file_path = './loss_of_sim_card.docx'  
    docx_file_path1 = './FG_contract_Tender.docx'  
    
    if reason == 'SIM Retrieval':
        doc = Document(docx_file_path)
    elif reason == "FG_contract_Tender":
        doc = Document(docx_file_path1)

    # Replace placeholders with user-provided values
    for paragraph in doc.paragraphs:
        if '[Your Full Name],' in paragraph.text:
                paragraph.text = paragraph.text.replace('[Your Full Name],', f'{first_name} {middle_name} {last_name},')
        if '[Your Address]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Your Address]', f'{address},')

        if '[Phone Number]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Phone Number]', f'{mobile},')

        if '[Date of Loss]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Date of Loss]', f'{date_of_loss},')

        if '[Date ]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Date]', f'{date},')

        # if '[Signature]' in paragraph.text:
        #     paragraph.text = paragraph.text.replace('[Signature]', f'{signature},')

        if '[Male]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Male]', f'{gender},')

        if '[Company Address]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Company Address]', f'{company_address},')

        if '[Position]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Position]', f'{position},')

        if '[Religion]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Religion]', f'{religion},')

        if '[Seal $ Sign]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Seal $ Sign]', f'{seal_sign},')

        if '[Year]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Year]', "2024")

        if '[MTN network]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[MTN network]', f'{network},')

    # Save the modified DOCX file
    modified_docx_path = f'{first_name + last_name}\'s modified_affidavit.docx'  # Replace with your desired file name
    doc.save(modified_docx_path)

    # Send email with the modified DOCX file as an attachment
    send_email(email, f'{first_name} {last_name}\'s Affidavit for {reason}', 'Thank you for using Bioentrust to generate an e-affidavit.', modified_docx_path)

    return jsonify(message='Affidavit generated and email sent successfully')

def send_email(to_email, subject, message, attachment_path):
    # Set up the SMTP server
    smtp_server = 'smtp.gmail.com'
    smtp_port = 587 
    smtp_username = 'mlsayabatech@gmail.com'
    smtp_password = 'hljx rioc tqwy cmca' # os.environ.get('SMTP_PASSWORD')

    # Create the email message
    msg = MIMEMultipart()
    msg['From'] = smtp_username
    msg['To'] = to_email
    msg['Subject'] = subject
    msg.attach(MIMEText(message, 'plain'))

    # Attach the modified DOCX file
    with open(attachment_path, 'rb') as file:
        attachment = MIMEApplication(file.read(), _subtype="docx")
        attachment.add_header('Content-Disposition', 'attachment', filename='Affidavit.docx')
        msg.attach(attachment)

    # Connect to the SMTP server and send the email
    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(smtp_username, smtp_password)
        server.sendmail(smtp_username, to_email, msg.as_string())

if __name__ == '__main__':
    app.run(debug=True)
